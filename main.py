import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import io
import base64

# Libraries for exporting to different formats my example is for word/docx file as a test
import xlsxwriter
import jinja2
import pdfkit
import docx
from docx.shared import Inches


class AdvancedDataAnalyzer:
    def __init__(self, df: pd.DataFrame):
        """
        Initialize the AdvancedDataAnalyzer with a DataFrame
        
        Parameters:
        -----------
        df : pandas.DataFrame
            Input dataframe to be analyzed
        """
        self.original_df = df.copy()
        self.df = df.copy()
        self.visualizations = {}  # Store visualizations as base64 images
    
    def _generate_visualization(self, plot_func):
        """
        Generate a plot and return it as a base64 encoded image.
        
        Parameters:
        -----------
        plot_func : callable
            Function to generate the plot
        
        Returns:
        --------
        base64 encoded image
        """
        buffer = io.BytesIO()  # Create an in-memory buffer
        plot_func()
        plt.savefig(buffer, format='png')
        plt.close()
        buffer.seek(0)
        return base64.b64encode(buffer.getvalue()).decode('utf-8')
    
    def find_missing_values(self):
        """
        Identify columns with missing values.
        
        Returns:
        --------
        DataFrame with missing value statistics
        """
        missing_info = pd.DataFrame({
            'Column': self.df.columns,
            'Missing Count': self.df.isnull().sum(),
            'Missing Percentage': (self.df.isnull().sum() / len(self.df)) * 100
        })
        return missing_info[missing_info['Missing Count'] > 0]
    
    def categorize_columns_by_type(self):
        """
        Categorize columns into different types.
        
        Returns:
        --------
        Dictionary with column type categories
        """
        return {
            'Numeric': list(self.df.select_dtypes(include=['int64', 'float64']).columns),
            'Categorical': list(self.df.select_dtypes(include=['object', 'category']).columns),
            'DateTime': list(self.df.select_dtypes(include=['datetime64']).columns),
            'Boolean': list(self.df.select_dtypes(include=['bool']).columns)
        }
    
    def handle_duplicates(self, remove=False):
        """
        Identify duplicate rows and optionally remove them.
        
        Parameters:
        -----------
        remove : bool, optional
            Whether to remove duplicate rows
        
        Returns:
        --------
        DataFrame containing duplicate rows
        """
        duplicate_rows = self.df[self.df.duplicated(keep=False)]
        if remove:
            self.df.drop_duplicates(inplace=True)
        return duplicate_rows
    
    def handle_constant_columns(self, remove=False):
        """
        Identify and optionally remove columns with constant values.
        
        Parameters:
        -----------
        remove : bool, optional
            Whether to remove constant columns
        
        Returns:
        --------
        DataFrame with constant column information
        """
        constant_columns = [col for col in self.df.columns if self.df[col].nunique() <= 1]
        constant_info = pd.DataFrame({
            'Column': constant_columns,
            'Constant Value': [self.df[col].iloc[0] for col in constant_columns]
        })
        if remove:
            self.df.drop(columns=constant_columns, inplace=True)
        return constant_info
    
    def generate_visualizations(self):
        """
        Generate visualizations as base64 encoded images.
        
        Returns:
        --------
        Dictionary of visualizations
        """
        visualizations = {}
        
        # Boxplot for outliers
        def boxplot():
            numeric_cols = self.df.select_dtypes(include=['int64', 'float64']).columns
            plt.figure(figsize=(15, 6))
            self.df[numeric_cols].boxplot()
            plt.title('Boxplot of Numeric Columns')
            plt.xticks(rotation=45)
            plt.tight_layout()
        visualizations['boxplot'] = self._generate_visualization(boxplot)
        
        # Distribution plots
        def distribution_plots():
            plot_columns = self.df.select_dtypes(include=['int64', 'float64', 'object']).columns[:6]
            fig, axes = plt.subplots(2, 3, figsize=(15, 10))
            axes = axes.ravel()
            for i, col in enumerate(plot_columns):
                if np.issubdtype(self.df[col].dtype, np.number):
                    sns.histplot(self.df[col], kde=True, ax=axes[i])
                else:
                    sns.countplot(x=col, data=self.df, ax=axes[i])
                axes[i].set_title(f'Distribution of {col}')
                axes[i].set_xlabel(col)
            plt.tight_layout()
        visualizations['distributions'] = self._generate_visualization(distribution_plots)
        
        return visualizations
    
    def export_to_word(self, output_path='data_analysis_report.docx'):
        """
        Export the analysis to a Word document.
        
        Parameters:
        -----------
        output_path : str
            Path to save the Word document
        """
        doc = docx.Document()
        doc.add_heading('Data Analysis Report', 0)
        
        # Missing values
        doc.add_heading('Missing Values', level=1)
        missing_values = self.find_missing_values()
        if not missing_values.empty:
            table = doc.add_table(rows=missing_values.shape[0]+1, cols=missing_values.shape[1])
            for col in range(missing_values.shape[1]):
                table.cell(0, col).text = missing_values.columns[col]
            for row in range(missing_values.shape[0]):
                for col in range(missing_values.shape[1]):
                    table.cell(row+1, col).text = str(missing_values.iloc[row, col])
        else:
            doc.add_paragraph("No missing values found.")
        
        # Add visualizations
        doc.add_heading('Visualizations', level=1)
        visualizations = self.generate_visualizations()
        
        # Save visualizations as temporary images
        boxplot_path = 'boxplot_temp.png'
        dist_path = 'distribution_temp.png'
        with open(boxplot_path, 'wb') as f:
            f.write(base64.b64decode(visualizations['boxplot']))
        with open(dist_path, 'wb') as f:
            f.write(base64.b64decode(visualizations['distributions']))
        
        # Add images to the document
        doc.add_picture(boxplot_path, width=Inches(6))
        doc.add_picture(dist_path, width=Inches(6))
        
        # Save the document
        doc.save(output_path)


def generate_sample_data():
    """
    Generate a sample dataset for testing the analyzer.
    
    Returns:
    --------
    pandas.DataFrame
        Sample dataset
    """
    np.random.seed(42)
    data = {
        "NumericColumn1": np.random.randint(0, 100, 50).astype(float),  # Convert to float for NaN
        "NumericColumn2": np.random.uniform(10.5, 75.3, 50),
        "CategoricalColumn": np.random.choice(['A', 'B', 'C'], 50),
        "BooleanColumn": np.random.choice([True, False], 50),
        "ConstantColumn": ['ConstantValue'] * 50,
        "DateColumn": pd.date_range(start="2023-01-01", periods=50, freq='D'),
    }
    data['NumericColumn1'][np.random.choice(50, 5, replace=False)] = np.nan  # Add NaN values
    data['CategoricalColumn'][np.random.choice(50, 3, replace=False)] = None  # Add None values
    return pd.DataFrame(data)


# Main execution for testing
if __name__ == "__main__":
    df = generate_sample_data() 
    analyzer = AdvancedDataAnalyzer(df)
    analyzer.export_to_word('data_analysis_report.docx')
    print("Data analysis report saved to data_analysis_report.docx")
